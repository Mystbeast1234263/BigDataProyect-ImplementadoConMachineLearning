#!/usr/bin/env python3
"""
Script para generar documento Word (.docx) desde INFORME_ML.md con formato APA profesional
"""

import os
import sys
import re
from pathlib import Path
from datetime import datetime

try:
    import markdown
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from markdown.extensions import tables
except ImportError as e:
    print(f"❌ Error: Faltan dependencias. Instala con:")
    print(f"   pip install markdown python-docx")
    print(f"\nError específico: {e}")
    sys.exit(1)


def add_page_number(doc):
    """Agrega número de página en el encabezado según formato APA"""
    section = doc.sections[0]
    header = section.header
    
    # Crear párrafo para el número de página
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Agregar número de página
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    
    # Formato del número de página
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)


def setup_apa_styles(doc):
    """Configura estilos APA en el documento"""
    
    # Configurar sección (márgenes APA: 1 pulgada = 2.54 cm)
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Estilo Normal (Times New Roman 12pt, doble espacio)
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(12)
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal_style.paragraph_format.space_after = Pt(0)
    
    # Estilo Título 1 (Centrado, Negrita, 12pt)
    if 'Heading 1' in doc.styles:
        h1_style = doc.styles['Heading 1']
        h1_font = h1_style.font
        h1_font.name = 'Times New Roman'
        h1_font.size = Pt(12)
        h1_font.bold = True
        h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h1_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        h1_style.paragraph_format.space_before = Pt(12)
        h1_style.paragraph_format.space_after = Pt(12)
    
    # Estilo Título 2 (Alineado izquierda, Negrita, 12pt)
    if 'Heading 2' in doc.styles:
        h2_style = doc.styles['Heading 2']
        h2_font = h2_style.font
        h2_font.name = 'Times New Roman'
        h2_font.size = Pt(12)
        h2_font.bold = True
        h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h2_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        h2_style.paragraph_format.space_before = Pt(12)
        h2_style.paragraph_format.space_after = Pt(6)
    
    # Estilo Título 3 (Alineado izquierda, Negrita, Cursiva, 12pt)
    if 'Heading 3' in doc.styles:
        h3_style = doc.styles['Heading 3']
        h3_font = h3_style.font
        h3_font.name = 'Times New Roman'
        h3_font.size = Pt(12)
        h3_font.bold = True
        h3_font.italic = True
        h3_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h3_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        h3_style.paragraph_format.space_before = Pt(12)
        h3_style.paragraph_format.space_after = Pt(6)
    
    # Estilo Título 4 (Alineado izquierda, Negrita, 12pt, con sangría)
    if 'Heading 4' in doc.styles:
        h4_style = doc.styles['Heading 4']
        h4_font = h4_style.font
        h4_font.name = 'Times New Roman'
        h4_font.size = Pt(12)
        h4_font.bold = True
        h4_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h4_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        h4_style.paragraph_format.left_indent = Inches(0.5)
        h4_style.paragraph_format.space_before = Pt(12)
        h4_style.paragraph_format.space_after = Pt(6)


def add_title_page(doc, title, authors, course, date):
    """Agrega portada con formato APA"""
    
    # Título (centrado, negrita, mayúsculas)
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title.upper())
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(12)
    title_run.bold = True
    title_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    title_para.paragraph_format.space_after = Pt(12)
    
    # Espacio
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Información del curso (centrado)
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(course)
    info_run.font.name = 'Times New Roman'
    info_run.font.size = Pt(12)
    info_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    info_para.paragraph_format.space_after = Pt(12)
    
    # Autores (centrado)
    for author in authors:
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_para.add_run(author)
        author_run.font.name = 'Times New Roman'
        author_run.font.size = Pt(12)
        author_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        author_para.paragraph_format.space_after = Pt(12)
    
    # Fecha (centrado)
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(date)
    date_run.font.name = 'Times New Roman'
    date_run.font.size = Pt(12)
    date_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    # Salto de página
    doc.add_page_break()


def parse_markdown_to_docx(md_content, doc):
    """Convierte contenido Markdown a elementos de Word"""
    
    # Convertir Markdown a HTML primero
    md = markdown.Markdown(extensions=['extra', 'codehilite', 'tables', 'fenced_code'])
    html = md.convert(md_content)
    
    # Parsear HTML manualmente (simplificado)
    lines = md_content.split('\n')
    in_code_block = False
    code_block_lines = []
    in_list = False
    list_level = 0
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Saltar líneas vacías
        if not line:
            i += 1
            continue
        
        # Bloques de código
        if line.startswith('```'):
            if in_code_block:
                # Cerrar bloque de código
                code_para = doc.add_paragraph()
                code_run = code_para.add_run('\n'.join(code_block_lines))
                code_run.font.name = 'Courier New'
                code_run.font.size = Pt(10)
                code_para.paragraph_format.left_indent = Inches(0.5)
                code_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                code_block_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(lines[i])
            i += 1
            continue
        
        # Títulos
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            title_text = line.lstrip('#').strip()
            
            if level == 1:
                para = doc.add_heading(title_text, level=1)
            elif level == 2:
                para = doc.add_heading(title_text, level=2)
            elif level == 3:
                para = doc.add_heading(title_text, level=3)
            elif level == 4:
                para = doc.add_heading(title_text, level=4)
            else:
                para = doc.add_paragraph(title_text)
                para.style = 'Heading 4'
            
            i += 1
            continue
        
        # Listas
        if line.startswith('- ') or line.startswith('* ') or re.match(r'^\d+\.', line):
            list_text = re.sub(r'^[-*]\s+|\d+\.\s+', '', line)
            
            para = doc.add_paragraph(style='List Bullet' if line.startswith(('- ', '* ')) else 'List Number')
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            
            # Procesar negritas en listas
            parts = re.split(r'(\*\*.*?\*\*)', list_text)
            for part in parts:
                if not part:
                    continue
                if part.startswith('**') and part.endswith('**'):
                    bold_text = part[2:-2]
                    run = para.add_run(bold_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.bold = True
                else:
                    run = para.add_run(part)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            
            i += 1
            continue
        
        # Tablas (básico)
        if '|' in line and not line.startswith('|---'):
            # Es una fila de tabla
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if not hasattr(doc, '_current_table'):
                doc._current_table = doc.add_table(rows=1, cols=len(cells))
                doc._current_table.style = 'Light Grid Accent 1'
                row = doc._current_table.rows[0]
                for j, cell_text in enumerate(cells):
                    cell = row.cells[j]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
            else:
                row = doc._current_table.add_row()
                for j, cell_text in enumerate(cells):
                    if j < len(row.cells):
                        row.cells[j].text = cell_text
                        for paragraph in row.cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(12)
            i += 1
            continue
        else:
            if hasattr(doc, '_current_table'):
                delattr(doc, '_current_table')
        
        # Separadores
        if line.startswith('---'):
            i += 1
            continue
        
        # Párrafos normales
        para = doc.add_paragraph()
        
        # Procesar texto preservando negritas y cursivas
        text = line
        para_format = para.paragraph_format
        para_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        
        # Procesar negritas (**texto**)
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                # Texto en negrita
                bold_text = part[2:-2]
                run = para.add_run(bold_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                # Texto en cursiva
                italic_text = part[1:-1]
                run = para.add_run(italic_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                # Código inline
                code_text = part[1:-1]
                run = para.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(11)
            else:
                # Texto normal
                run = para.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        
        i += 1


def generar_word_informe_ml():
    """Genera documento Word desde INFORME_ML.md con formato APA"""
    
    # Rutas
    base_dir = Path(__file__).parent
    md_file = base_dir / "docs" / "INFORME_ML.md"
    output_docx = base_dir / "docs" / "INFORME_ML_APA.docx"
    
    # Verificar que existe el archivo Markdown
    if not md_file.exists():
        print(f"❌ Error: No se encontró el archivo {md_file}")
        sys.exit(1)
    
    print(f"📖 Leyendo archivo: {md_file}")
    
    # Leer contenido Markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        markdown_content = f.read()
    
    # Crear documento Word
    print("📝 Creando documento Word...")
    doc = Document()
    
    # Configurar estilos APA
    setup_apa_styles(doc)
    
    # Agregar número de página
    add_page_number(doc)
    
    # Extraer información para portada
    title = "INFORME TÉCNICO - PRÁCTICA N°4: FUNDAMENTOS DE MACHINE LEARNING"
    authors = [
        "Joshua Chavez Abirari - Desarrollador principal",
        "Joel Israel Lopez Ticlla - Compañero de equipo (Apoyo y colaboración)"
    ]
    course = "ASIGNATURA: TECNOLOGÍAS EMERGENTES I\nPROYECTO: GAMC Big Data Dashboard - Sistema de Análisis y Predicción con Machine Learning"
    date = "Diciembre 2024"
    
    # Agregar portada
    add_title_page(doc, title, authors, course, date)
    
    # Procesar contenido (saltar encabezado y metadatos)
    lines = markdown_content.split('\n')
    start_idx = 0
    
    # Buscar donde empieza el contenido real (después de los metadatos)
    for i, line in enumerate(lines):
        if line.startswith('## ÍNDICE') or line.startswith('## 1.'):
            start_idx = i
            break
    
    content_lines = '\n'.join(lines[start_idx:])
    
    # Convertir Markdown a Word
    print("🔄 Convirtiendo Markdown a Word...")
    parse_markdown_to_docx(content_lines, doc)
    
    # Guardar documento
    print("💾 Guardando documento Word...")
    try:
        doc.save(output_docx)
        print(f"✅ Documento Word generado exitosamente: {output_docx}")
        print(f"   Tamaño del archivo: {output_docx.stat().st_size / 1024:.2f} KB")
        return True
    except Exception as e:
        print(f"❌ Error al guardar documento: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == "__main__":
    print("=" * 60)
    print("🚀 Generador de Word - Informe ML con Formato APA")
    print("=" * 60)
    print()
    
    success = generar_word_informe_ml()
    
    print()
    if success:
        print("✨ ¡Proceso completado exitosamente!")
        print("\n📋 Características del documento:")
        print("   ✓ Formato APA (márgenes 1 pulgada, Times New Roman 12pt)")
        print("   ✓ Interlineado doble")
        print("   ✓ Portada con formato académico")
        print("   ✓ Numeración de páginas")
        print("   ✓ Estilos de títulos según APA")
    else:
        print("⚠️  El proceso falló. Revisa los errores arriba.")
        sys.exit(1)

